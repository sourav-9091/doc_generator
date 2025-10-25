import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_formatted_paragraph(doc, text, style=None, font_size=12):
    """
    Add a paragraph with inline markdown formatting (**bold**, *italic*) and font size.
    """
    p = doc.add_paragraph(style=style)
    # Split by bold (**text**) or italic (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(font_size)
        else:
            run = p.add_run(part)
            run.font.size = Pt(font_size)
    return p

def create_word_doc(title, prepared_by, content, images=[], output_path="SAP_Documentation.docx"):
    doc = Document()

    # Define main headings with numbering
    main_headings = [
        "Purpose",
        "Scope",
        "Background",
        "Root Cause Analysis",
        "Design Solution",
        "Objects Changed"
    ]

    heading_color = RGBColor(0, 51, 102)  # Dark blue color

    # Technical Specification heading first
    h1 = doc.add_heading("Technical Specification", level=1)
    run = h1.runs[0]
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = heading_color

    doc.add_paragraph(f"Title: {title}", style=None).runs[0].font.size = Pt(12)
    doc.add_paragraph(f"Prepared By: {prepared_by}", style=None).runs[0].font.size = Pt(12)

    current_table = None

    # Number map for headings
    heading_numbers = {h: f"{i+1}. " for i, h in enumerate(main_headings)}

    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue

        # Check if line matches any main heading
        matched_heading = None
        for heading in main_headings:
            if heading.lower() in line.lower():
                matched_heading = heading
                break

        if matched_heading:
            h_text = heading_numbers[matched_heading] + matched_heading
            h = doc.add_heading(h_text, level=1)
            run = h.runs[0]
            run.font.size = Pt(26)
            run.font.bold = True
            run.font.color.rgb = heading_color
            current_table = None
            continue

        # Bullet list
        if line.startswith("* "):
            add_formatted_paragraph(doc, line[2:], style='List Bullet', font_size=12)
            current_table = None
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', line):
            item_text = re.sub(r'^\d+\.\s+', '', line)
            add_formatted_paragraph(doc, item_text, style='List Number', font_size=12)
            current_table = None
            continue

        # Tables (markdown style)
        if line.startswith("|") and "|" in line:
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if current_table is None:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, cell in enumerate(cells):
                    hdr_cells[i].text = cell
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
                current_table = table
            else:
                row_cells = current_table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
            continue

        # Normal paragraph with inline formatting
        add_formatted_paragraph(doc, line, font_size=12)
        current_table = None

    # Add images
    for img_path in images:
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5))

    doc.save(output_path)
    return output_path
